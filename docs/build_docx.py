# -*- coding: utf-8 -*-
"""Build a professional .docx reference document for the ITM GOI Examination Portal."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MAROON      = RGBColor(0x8A, 0x1B, 0x2C)
MAROON_DEEP = RGBColor(0x4A, 0x0D, 0x17)
GOLD        = RGBColor(0xA9, 0x82, 0x2F)
INK         = RGBColor(0x24, 0x1A, 0x1C)
SLATE       = RGBColor(0x5F, 0x57, 0x5A)

OUT = r"E:\ASSESMENT\docs\ITM-GOI-Examination-Portal.docx"

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for i, (nm, sz, serif) in enumerate([("Heading 1", 16, True), ("Heading 2", 12.5, False), ("Heading 3", 11, False)]):
    st = doc.styles[nm]
    st.font.name = "Georgia" if serif else "Calibri"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = MAROON if i == 0 else INK
    st.paragraph_format.space_before = Pt(18 if i == 0 else 12)
    st.paragraph_format.space_after = Pt(6 if i == 0 else 4)
    st.paragraph_format.keep_with_next = True

# page margins
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1.0)

# ---------- helpers ----------
def shade(el, hexcolor):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    el.append(sh)

def para_border_bottom(p, hexcolor="A9822F", sz=8):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), hexcolor)
    pbdr.append(bottom); pPr.append(pbdr)

def h1(clause, text):
    p = doc.add_heading(level=1)
    r = p.add_run(clause + "\u2003")
    r.font.color.rgb = GOLD; r.font.name = "Consolas"; r.font.size = Pt(13)
    r2 = p.add_run(text)
    r2.font.color.rgb = MAROON; r2.font.name = "Georgia"; r2.font.size = Pt(16); r2.bold = True
    para_border_bottom(p)
    return p

def h2(text): return doc.add_heading(text, level=2)
def h3(text): return doc.add_heading(text, level=3)

def body(text, style=None):
    p = doc.add_paragraph(style=style)
    _rich(p, text)
    return p

def _rich(p, text):
    """*bold*, `mono` inline markup."""
    import re
    for tok in re.split(r'(\*[^*]+\*|`[^`]+`)', text):
        if not tok:
            continue
        if tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = MAROON
        else:
            p.add_run(tok)

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _rich(p, it)

def numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        _rich(p, it)

def note(lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    shade(p._p.get_or_add_pPr(), "F7EEF0")
    r = p.add_run(lead + " "); r.bold = True; r.font.color.rgb = MAROON; r.font.size = Pt(9.5)
    _rich(p, text)
    for rr in p.runs[1:]:
        rr.font.size = Pt(9.5)

def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    shade(p._p.get_or_add_pPr(), "F1E9E3")
    for i, ln in enumerate(lines):
        if i: p.add_run().add_break()
        r = p.add_run(ln); r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = INK

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        pr = hdr[i].paragraphs[0]; run = pr.add_run(htext.upper())
        run.bold = True; run.font.size = Pt(8); run.font.color.rgb = SLATE
        shade(hdr[i]._tc.get_or_add_tcPr(), "EFE6E1")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            _rich(cells[i].paragraphs[0], val)
            for rr in cells[i].paragraphs[0].runs:
                rr.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    it = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = "Right-click and choose \u201cUpdate field\u201d to build the table of contents."
    it.append(t); fld.append(it)
    p._p.append(fld)

def footer_text():
    for s in doc.sections:
        f = s.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = f.add_run("ITM Group of Institutions, Gwalior  \u00b7  Online Examination Portal  \u00b7  Page ")
        r.font.size = Pt(8); r.font.color.rgb = SLATE
        fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
        rr = OxmlElement("w:r"); tt = OxmlElement("w:t"); tt.text = "1"; rr.append(tt); fld.append(rr)
        f._p.append(fld)

# =====================================================================
#  COVER
# =====================================================================
k = doc.add_paragraph()
kr = k.add_run("ITM GROUP OF INSTITUTIONS  \u00b7  GWALIOR")
kr.bold = True; kr.font.size = Pt(9); kr.font.color.rgb = GOLD
k.paragraph_format.space_after = Pt(2)

title = doc.add_paragraph()
tr = title.add_run("Online Examination Portal")
tr.font.name = "Georgia"; tr.font.size = Pt(30); tr.bold = True; tr.font.color.rgb = MAROON
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sr = sub.add_run("Platform Reference Document")
sr.font.size = Pt(13); sr.font.color.rgb = SLATE
para_border_bottom(sub, "A9822F", 12)
sub.paragraph_format.space_after = Pt(14)

body("A self-hosted, AI-proctored assessment platform built on Moodle 4.5 LTS \u2014 with a custom "
     "exam-cell toolkit, Safe Exam Browser kiosk lockdown, live invigilation, and the institute's "
     "own visual identity. Runs from a single Windows machine, with no cloud hosting and no Docker.")

doc.add_paragraph().paragraph_format.space_after = Pt(4)
table(["Field", "Value"],
      [["Platform", "Moodle 4.5 LTS"],
       ["Technology stack", "Apache 2.4  \u00b7  PHP 8.3 (mod_php)  \u00b7  MariaDB 11.4 LTS"],
       ["Runtime", "Native Windows services \u2014 no containers"],
       ["Portal address", "`http://localhost:8080`"],
       ["Repository", "github.com/Neerajparashar1/ASSESMENT"],
       ["Document", "Reference / capability overview"]],
      widths=[1.7, 4.6])

doc.add_page_break()

# =====================================================================
#  CONTENTS
# =====================================================================
ct = doc.add_paragraph()
cr = ct.add_run("Contents"); cr.bold = True; cr.font.size = Pt(13); cr.font.color.rgb = MAROON
add_toc()
doc.add_page_break()

# =====================================================================
#  §1 OVERVIEW
# =====================================================================
h1("\u00a71", "Overview")
body("The portal lets the examination cell run secured, invigilated online exams for students \u2014 "
     "creating question banks, building exams, locking the candidate's machine with Safe Exam "
     "Browser, watching every attempt live, and exporting results \u2014 all from one server on campus.")
body("It is a standard *Moodle 4.5 LTS* installation, hardened for exams and extended with two "
     "purpose-built plugins:")
bullets([
    "*local_examwizard* \u2014 the exam cell's control centre: question uploader, a four-step exam "
    "builder, the Exam Control hub, live exam control, results export, bulk student onboarding, and "
    "the SEB quit-password manager.",
    "*local_sebkiosk* \u2014 makes leaving Safe Exam Browser finalise the attempt, and carries the "
    "front-end theme runtime.",
])
body("Contributed modules add webcam proctoring (`quizaccess_proctoring`), the SEB access rule "
     "(`quizaccess_seb`), and the Boost Union theme, skinned in the ITM Group of Institutions identity.")
h3("Design intent")
body("Every exam is one attempt, one question per page, shuffled, time-limited, copy-protected, and "
     "\u2014 where required \u2014 only openable inside a locked-down browser with the webcam on. If a "
     "candidate's machine fails or they leave the browser, the attempt is submitted automatically and "
     "cannot be resumed.")

# =====================================================================
#  §2 ARCHITECTURE
# =====================================================================
h1("\u00a72", "Architecture & technology")
body("The platform runs directly on Windows as three services plus a scheduled task \u2014 no "
     "containers. Docker was evaluated but its Linux engine does not start on the host machine, so "
     "the stack was ported to run natively.")
table(["Layer", "Component", "Notes"],
      [["Application", "Moodle 4.5 LTS (MOODLE_405_STABLE)", "Web root at `native/stack/moodle`"],
       ["Web server", "Apache 2.4 with mod_php", "Service `eap-apache`  \u00b7  port 8080"],
       ["Language", "PHP 8.3, Thread-Safe", "OPcache, 512 MB limit, 100 MB uploads"],
       ["Database", "MariaDB 11.4 LTS", "Service `eap-mariadb`  \u00b7  utf8mb4_unicode_ci"],
       ["Scheduler", "Windows Task EAP-Moodle-Cron", "Runs Moodle cron every minute as SYSTEM"],
       ["Theme", "Boost Union + custom Raw SCSS", "Compiled from `config/moodle/custom.scss`"]],
      widths=[1.15, 2.6, 2.55])
body("Application data lives at `E:\\ASSESMENT\\moodledata`; all secrets and site settings are read "
     "from `E:\\ASSESMENT\\.env`. The database data directory and downloadable installers are "
     "excluded from source control.")

# =====================================================================
#  §3 SETUP
# =====================================================================
h1("\u00a73", "Setup & daily operation (no Docker)")
body("The platform is a self-contained folder. Everything it needs \u2014 PHP, Apache, MariaDB and "
     "Moodle \u2014 travels inside the repository, so setup is a `git clone` plus one script. There is "
     "no Docker, no separate installers, and no download step during setup. The script derives every "
     "path from the folder it sits in, so the project can live on any drive or directory.")

h2("3.1  What you need")
bullets([
    "Windows 10 (build 1809 or later) or Windows 11, 64-bit.",
    "An account with *local administrator* rights \u2014 the setup script self-elevates via UAC.",
    "*Windows PowerShell 5.1* (ships with Windows); PowerShell 7 also works.",
    "*Git for Windows*.",
    "About *2 GB* of free disk space \u2014 the repository carries the whole server stack.",
    "Port *8080* free, or choose another with `-Port`.",
])

h2("3.2  Get the code")
code_block(["git config --global core.longpaths true",
            "git clone https://github.com/Neerajparashar1/ASSESMENT.git",
            "cd ASSESMENT"])
body("It can be cloned anywhere \u2014 `C:\\exam`, `D:\\ASSESMENT`, an external drive. *Avoid* a "
     "OneDrive-synced folder or a very deep path; file-locking and the 260-character path limit "
     "cause confusing failures. A short path such as `C:\\exam\\ASSESMENT` is safest.")

h2("3.3  Exclude the folder from antivirus")
body("Windows Defender and other antivirus tools occasionally quarantine part of the bundled server "
     "(a MariaDB or PHP executable) during the clone or during setup, which then fails in ways that "
     "are hard to diagnose. Before running setup, add the project folder as an exclusion:")
body("Windows Security \u2192 Virus & threat protection \u2192 Manage settings \u2192 Exclusions "
     "\u2192 Add an exclusion \u2192 Folder \u2192 select the ASSESMENT folder.")
body("If something was already quarantined, restore it from Windows Security \u2192 Protection "
     "history, add the exclusion, and clone again.")

h2("3.4  Create the settings file")
body("`.env` holds every secret and is deliberately not in the repository. Copy the template and "
     "edit it:")
code_block(["copy .env.example .env",
            "notepad .env"])
table(["Key", "Set to"],
      [["`HTTP_PORT`", "8080, or your chosen port"],
       ["`MOODLE_WWWROOT`", "`http://localhost:8080` \u2014 must match the port"],
       ["`MOODLE_ADMIN_USER`", "the site administrator login"],
       ["`MOODLE_ADMIN_PASS`", "a strong password: 8+ chars, upper + lower + digit + symbol"],
       ["`MOODLE_ADMIN_EMAIL`", "a real address"],
       ["`DB_ROOT_PASS`", "a new, strong MariaDB root password"],
       ["`DB_PASS`", "a new, strong password for the `moodle` database user"],
       ["`SEB_QUIT_PASSWORD`", "the Safe Exam Browser quit password"],
       ["`MOODLE_SITE_FULLNAME` / `_SHORTNAME`", "shown across the portal"],
       ["`INNODB_BUFFER_POOL_SIZE`", "`1G` on a 4 GB machine; `2G` or more on 8 GB+"],
       ["`TZ`", "your timezone, e.g. `Asia/Kolkata`"]],
      widths=[2.5, 3.8])
body("Leave `DB_HOST`, `DB_NAME`, `DB_USER` and `MOODLE_BRANCH` at their defaults.")

h2("3.5  Run the installer")
body("Open PowerShell *in the project folder* and run:")
code_block(["powershell -ExecutionPolicy Bypass -File .\\native\\Setup-MoodleNative.ps1 -Reinstall"])
note("On a different PC \u2014",
     "always include `-Reinstall`. The repository carries a `config.php` and generated Apache / "
     "MariaDB configuration files with the *original* machine's paths and a machine-specific "
     "database. `-Reinstall` drops those and rebuilds every path, the database and the config for "
     "*this* machine. On the machine the repository was built on, run it *without* `-Reinstall`.")
body("The script performs, in order:")
numbered([
    "Installs the Visual C++ 2015\u20132022 runtime if it is missing.",
    "Verifies the bundled PHP 8.3, Apache 2.4, MariaDB 11.4 and Moodle 4.5 are present "
    "(they ship in the repo; it only downloads if one is missing).",
    "Generates `php.ini`, the Apache virtual host and `my.ini` \u2014 with *this machine's* paths.",
    "Initialises MariaDB, registers the `eap-mariadb` service, creates the `moodle` database and its user.",
    "Runs Moodle's non-interactive installer against `<project>\\moodledata`.",
    "Installs the contrib plugins \u2014 Boost Union theme, proctoring, VPL.",
    "Writes the exam-hardening block into `config.php` (anti-cheat footer, autosave, proxied-SSL "
    "handling, the path to PHP).",
    "Runs `post_install.php` \u2014 theme and custom SCSS, SEB and proctoring defaults, quiz "
    "hardening, the Invigilator and Exam Hall Invigilator roles.",
    "Registers the `eap-apache` service, the one-minute `EAP-Moodle-Cron` task and a firewall rule "
    "for the port.",
    "Generates a starter `.seb` file from `.env`.",
])
body("It prints the portal URL and the admin login when done, and is safe to re-run.")
note("If you move the folder later \u2014",
     "the generated Apache and MariaDB configs still point at the old location. Re-run "
     "`Setup-MoodleNative.ps1 -Reinstall` from the new location (this rebuilds the database), or "
     "hand-edit `my.ini`, `Apache24\\conf\\httpd.conf`, `Apache24\\conf\\extra\\httpd-eap.conf` and "
     "`native\\stack\\moodle\\config.php`, then re-register the services.")

h2("3.6  Apply the exam-specific configuration")
body("These create the kiosk SEB template, ensure both invigilation roles, and compile the ITM "
     "theme. A `-Reinstall` already covers the roles and theme through `post_install.php`; run the "
     "SEB template script on every fresh machine, and re-run any of them after a related change.")
code_block(["set PHP=native\\stack\\php\\php.exe",
            "%PHP% native\\Setup-SebLockdownTemplate.php --all",
            "%PHP% native\\Setup-InvigilatorRoles.php",
            "%PHP% native\\apply_itm_theme.php"])

h2("3.7  Verify the installation")
bullets([
    "`native\\Manage-Moodle.ps1 status` \u2014 both services `Running`, the cron task present, and an "
    "`HTTP :8080  200 OK` line.",
    "Open `http://localhost:8080` \u2014 the branded split-view login page appears.",
    "Sign in as the admin. *Site administration \u2192 Notifications* should report no problems.",
    "*Site administration \u2192 Server \u2192 Environment* \u2014 every row should pass.",
])

h2("3.8  Daily operation")
body("Both services are set to start automatically, so after a reboot the portal is already "
     "running. Double-click `start.bat` to bring the stack up manually and `stop.bat` to take it "
     "down. The PowerShell wrapper `native\\Manage-Moodle.ps1` covers the rest:")
table(["Command", "Action"],
      [["`status`", "Service state and an HTTP probe"],
       ["`start` / `stop` / `restart`", "Control the two services (needs an elevated shell)"],
       ["`cron`", "Run one Moodle cron cycle now"],
       ["`purge`", "Purge all Moodle caches"],
       ["`upgrade`", "Run the Moodle upgrade after adding or updating plugins"],
       ["`import <csv>`", "Bulk-import students from a roster"],
       ["`backup`", "Database dump + moodledata archive to `backups\\`"]],
      widths=[2.0, 4.3])
body("Grade export to ODS, XLS, TXT and XML is enabled site-wide from the gradebook.")

h2("3.9  Useful installer switches")
bullets([
    "`-Reinstall` \u2014 drop and rebuild the database, config and data directory (required on a new machine).",
    "`-Port 8888` \u2014 serve on a different port; also updates the portal URL.",
    "`-PhpSeries 8.2` \u2014 use PHP 8.2 instead of 8.3.",
    "`-SkipPlugins` \u2014 do not (re)install the contrib plugins.",
    "`-ApacheZip / -PhpZip / -MariaDbZip / -MoodleZip <path>` \u2014 use a hand-downloaded build "
    "if a bundled one is missing or blocked.",
    "`-InstallRoot / -DataRoot <path>` \u2014 place the stack or the data directory elsewhere.",
])

h2("3.10  Troubleshooting")
table(["Symptom", "Cause and fix"],
      [["Setup stops on the Apache download",
        "Only if the bundled binary is missing. Apache Lounge sits behind Cloudflare and sometimes "
        "blocks scripted downloads \u2014 fetch the VS17 `httpd` zip by hand and pass "
        "`-ApacheZip <path>`."],
       ["`php8apache2_4.dll` missing / Apache will not start",
        "A non-thread-safe PHP build was used. Re-run with `-PhpZip <thread-safe zip> -Reinstall`."],
       ["Port 8080 already in use",
        "`netstat -ano | findstr :8080` to find the owner; stop it, or run setup with `-Port 8888`."],
       ["`Error: Database connection failed`",
        "MariaDB did not start, or the passwords in `config.php` do not match `.env`. Check "
        "`native\\stack\\logs\\` and re-run with `-Reinstall`."],
       ["Blank page or HTTP 500",
        "Read `native\\stack\\logs\\php_error.log` and `native\\stack\\Apache24\\logs\\error.log`. "
        "Usually a missing PHP extension or a `config.php` carried over from another machine \u2014 "
        "re-run with `-Reinstall`."],
       ["\u201cThis server may be accessible only via \u2026\u201d",
        "`$CFG->wwwroot` points elsewhere \u2014 a demo tunnel (section 10) was left on. Run "
        "`demo-stop.ps1` (or `demo-cf-stop.bat`)."],
       ["Services will not start after a reboot",
        "An antivirus quarantined a stack binary. Restore it from Protection history, add the "
        "folder exclusion (3.3), then re-run with `-Reinstall`."],
       ["`git clone` fails with \u201cFilename too long\u201d",
        "`git config --global core.longpaths true`, then clone into a shorter path."]],
      widths=[2.35, 3.95])
body("If you would rather not use the script at all, `native\\README-NATIVE.md` section 4 has the "
     "full manual sequence.")

# =====================================================================
#  §4 EXAM WIZARD
# =====================================================================
h1("\u00a74", "Exam Wizard toolkit")
body("`local_examwizard` is a custom plugin built entirely on Moodle's supported APIs \u2014 the quiz "
     "engine and question bank are unmodified. Staff reach everything from one landing page, the "
     "*Exam Control* hub.")
h3("Question uploader")
body("Upload questions as CSV, XLSX, GIFT, XML or plain text. Column headers are matched "
     "synonym-tolerantly. Before anything is imported, a visual preview renders every question as a "
     "card \u2014 green when valid, red with the specific error when not, correct options highlighted, "
     "with summary counts. Supports single-answer MCQ, multiple-response, true/false and short "
     "answer; GIFT and XML are handed to Moodle's own importers. Optional negative marking sets wrong "
     "options on single-answer MCQs to minus one-third.")
h3("Four-step Create-Exam wizard")
body("Basics \u2192 Questions \u2192 Rules \u2192 Review. It applies exam-sensible defaults \u2014 a "
     "time limit, automatic submission when time expires, deferred feedback, free navigation, results "
     "released only after the exam closes \u2014 forces a single attempt, and wires up Safe Exam "
     "Browser and proctoring in the same flow.")
h3("Exam Control hub")
body("The operations home page:")
bullets([
    "At-a-glance counts \u2014 exams, students, subjects, and how many exams are live now.",
    "An auto-detected getting-started checklist (branding, students, an exam, the SEB password, a test run).",
    "Quick actions and a *Your exams* table with a status badge (live / open / scheduled / closed), an "
    "\u201cin progress / submitted\u201d count, and Monitor, Grades and \u201cGet .seb\u201d links per exam.",
    "A *Live-now spotlight* \u2014 for each running exam: in-progress, submitted and not-started counts "
    "with a progress bar; the page refreshes itself every 30 seconds while an exam is live.",
])
h3("Live exam control")
body("Steer an exam while it runs, without editing it:")
table(["Scope", "Actions"],
      [["Whole exam", "Pause (close now)  \u00b7  Reopen for two hours  \u00b7  End & submit everyone"],
       ["Per candidate", "End & submit now  \u00b7  Give +15 minutes (per-user override)  \u00b7  "
                         "Resume an abandoned attempt  \u00b7  Delete attempt (teachers / managers only)"]],
      widths=[1.4, 4.9])
h3("Results & export")
body("A per-candidate table \u2014 score, percentage, pass/fail, time taken, attempt state \u2014 with "
     "summary chips for attempted-vs-enrolled, average percentage and pass rate, and a one-click CSV "
     "download.")
h3("Bulk student onboarding")
body("Upload a roster (CSV or XLSX). A validated preview shows which rows will create a new user and "
     "which already exist; on confirm it creates the missing accounts and enrols every valid row as a "
     "student. A default password field fills blank cells.")
h3("SEB quit-password manager")
body("View or change the master Safe Exam Browser quit password, and optionally push it to every "
     "SEB-enabled quiz at once. The screen warns that changing it invalidates any `.seb` files "
     "already handed out.")

# =====================================================================
#  §5 SEB
# =====================================================================
h1("\u00a75", "Safe Exam Browser lockdown")
body("Exams that require it are opened only inside Safe Exam Browser, configured from the *EAP Kiosk "
     "Lockdown* template:")
bullets([
    "*Windows kiosk* \u2014 a fresh isolated desktop and the Explorer shell killed, so Task Manager, "
    "the Start menu and the taskbar are unreachable.",
    "*Every escape hotkey disabled* \u2014 Alt+Tab, Alt+F4, Alt+Esc, Ctrl+Esc, Esc, F1\u2013F12 and PrintScreen.",
    "No application or user switching; a single display only; screen mirroring, virtual machines and "
    "screen sharing blocked.",
    "The browser is URL-filtered to the portal host.",
    "A Quit button is present and password-protected \u2014 the one sanctioned way out.",
])
h3("Leaving the browser submits the exam")
body("The `local_sebkiosk` plugin turns any exit into a submission. The SEB Quit link and a "
     "page-unload beacon both call an endpoint that finalises every in-progress attempt for that "
     "student, so it can never be reopened.")
h3("Backstop for a hard kill")
body("If a machine is powered off or force-quit mid-exam, the time limit plus automatic overdue "
     "submission plus the one-minute cron task submit the attempt when its time expires \u2014 nothing "
     "is left hanging.")
h3("Not tunnel-compatible")
body("A live SEB attempt is bound to the server's real address and cannot run over a temporary demo "
     "tunnel (section 10). Tunnels are for showing the portal and admin, not for taking a locked exam.")

# =====================================================================
#  §6 PROCTORING
# =====================================================================
h1("\u00a76", "AI proctoring")
body("The `quizaccess_proctoring` module captures a webcam snapshot roughly every 30 seconds during "
     "an attempt, stores the images, and gives staff a proctoring review report alongside the quiz "
     "results. A face-present check runs when the candidate starts. No paid cloud face-matching "
     "service is used; review is by a person against the captured images.")

# =====================================================================
#  §7 ROLES
# =====================================================================
h1("\u00a77", "Roles & access control")
body("Access is least-privilege. The site administrator is a single named account; everything else "
     "is a role assigned per course, per activity, or site-wide.")
table(["Role", "Can", "Cannot"],
      [["Site administrator (`examadmin`)",
        "Everything \u2014 the only account with full platform and server settings control.", "\u2014"],
       ["Invigilator / Proctor (faculty)",
        "Build quizzes and question banks, manage SEB settings, run proctoring review, grade, run live exam control.",
        "Site administration."],
       ["Exam Hall Invigilator (hall staff)",
        "Watch the monitor report and the live proctoring feed; run live exam control \u2014 pause, reopen, extend a student, force-submit, resume.",
        "Create or edit exams, edit questions, grade, delete attempts, reach site administration."],
       ["Student", "Sit the exams assigned to them.", "Everything else."]],
      widths=[1.6, 2.75, 1.95])
body("The hall-invigilator role is backed by a dedicated capability, `local/examwizard:control`, "
     "deliberately kept separate from Moodle's `mod/quiz:manage` so a proctor can steer a running "
     "exam without ever gaining editing rights. Assigned at site level, one login can oversee every "
     "exam session at once.")

# =====================================================================
#  §8 INTEGRITY
# =====================================================================
h1("\u00a78", "Exam integrity & anti-cheat")
table(["Measure", "Detail"],
      [["One attempt", "Every exam allows a single attempt; once finished it cannot be restarted."],
       ["Shuffled, paged", "Questions are shuffled and shown one per page, with free navigation between them."],
       ["Clipboard blocked", "Copy, cut, paste, right-click and drag are disabled on the attempt page; typing into answer fields still works."],
       ["Continuous autosave", "Answers save every 15 seconds and on every page change, so nothing is lost to a disconnect."],
       ["Time-limit submission", "When the clock runs out the attempt is submitted automatically, with no grace period."],
       ["Config-key gating", "For SEB exams, the browser's configuration key is checked on every request \u2014 viewing, starting, each page and submitting \u2014 not only on quit."]],
      widths=[1.7, 4.6])

# =====================================================================
#  §9 INTERFACE
# =====================================================================
h1("\u00a79", "Interface & branding")
body("The portal wears the ITM Group of Institutions identity \u2014 deep maroon and academic gold, "
     "the Inter and Source Serif typefaces, the official logo, and the \u201cThink Big \u00b7 Think "
     "Beyond\u201d motto set in the institute's own kicker style.")
bullets([
    "*Split-view sign-in* \u2014 an auto-advancing campus slideshow with college information on the "
    "left, the login card on the right with assurance chips (locked-down browser, AI proctoring, "
    "instant result) and a support bar.",
    "*Home & dashboard* \u2014 a branded hero band with the logo, the motto, a \u201cGo to my "
    "exams\u201d call to action and a security badge over a campus slideshow; a campus photograph "
    "fills the page margins; content blocks are styled as clean cards.",
    "*Available courses* render as a responsive, colour-coded card grid with an initial-letter chip "
    "and a whole-card link.",
    "*Exam-taking screen* \u2014 a distraction-free shell, a top progress bar, a question-navigator "
    "colour legend, a floating countdown that turns amber under five minutes and pulses red under a "
    "minute, and whole-option-card selection.",
    "*Responsive* \u2014 the hero banners stack, the motto wraps, and hover-only cues drop on phones.",
])
body("All static asset links are root-relative, so the site renders identically whether it is served "
     "from `localhost` or a temporary public address. Styling has one source of truth, "
     "`config/moodle/custom.scss`, compiled into the theme by `native/apply_itm_theme.php`; "
     "behaviour lives in `local/sebkiosk/exam-ui.js`.")

# =====================================================================
#  §10 DEMO
# =====================================================================
h1("\u00a710", "Sharing a live demo")
body("The portal can be shown to someone off-campus over a temporary tunnel from the host machine, "
     "with no cloud hosting. Two paths are wired up:")
table(["Method", "Command", "Trade-off"],
      [["Cloudflare", "`demo-cf.bat`  /  `demo-cf-stop.bat`",
        "Double-click, no account, no admin rights. Fresh random URL each run; quick tunnels can be flaky."],
       ["ngrok", "`demo.ps1 -Domain <name>.ngrok-free.app`  /  `demo-stop.ps1`",
        "Needs a free ngrok account and a claimed static domain; the URL is then stable and reusable."]],
      widths=[0.95, 2.55, 2.8])
body("Both temporarily repoint the site at the public address and turn on proxied-SSL handling; the "
     "stop script restores `localhost`. Safe Exam Browser exams do not work over a tunnel (section 5).")

# =====================================================================
#  §11 SECURITY
# =====================================================================
h1("\u00a711", "Security posture & open items")
body("Indicative resistance to a determined candidate: on an institute-managed lab PC the setup is "
     "strong; on a student's own laptop it is good, with webcam proctoring carrying much of the "
     "weight in practice.")
h3("Open items")
table(["Item", "Status", "Detail"],
      [["Browser Exam Key check", "Decision pending",
        "Config-key verification is on, but the SEB application binary itself is not yet pinned. "
        "Adding the official build's key to each exam would close the \u201cpatched SEB\u201d gap."],
       ["Kiosk template smoke test", "Pending",
        "The EAP Kiosk Lockdown template needs one manual launch on the institute's SEB build to "
        "confirm the stricter keys are accepted."],
       ["MariaDB loader stub", "Cosmetic",
        "A 13 KB stub file is in a delete-pending state. The running service uses a sibling binary "
        "and is unaffected; restore the file after a reboot."],
       ["Front-end polish backlog", "Scoped, not started",
        "Tables, forms, alerts, buttons, breadcrumbs and empty states have a queued modernisation pass."]],
      widths=[1.5, 1.25, 3.55])

# =====================================================================
#  §12 REFERENCE
# =====================================================================
h1("\u00a712", "Reference \u2014 files & accounts")
h3("Scripts & source of truth")
table(["Path", "Purpose"],
      [["`start.bat` / `stop.bat`", "Run or stop the stack"],
       ["`demo-cf.bat` / `demo-cf-stop.bat`", "Cloudflare demo tunnel"],
       ["`demo.ps1` / `demo-stop.ps1`", "ngrok / manual demo tunnel"],
       ["`native/Setup-MoodleNative.ps1`", "Full first-time install"],
       ["`native/Manage-Moodle.ps1`", "Day-to-day operations wrapper"],
       ["`native/Setup-SebLockdownTemplate.php`", "Create / apply the kiosk SEB template"],
       ["`native/Setup-InvigilatorRoles.php`", "Create the two invigilation roles"],
       ["`native/Import-Students.ps1`", "Bulk student import from the command line"],
       ["`native/apply_itm_theme.php`", "Compile custom.scss into the theme"],
       ["`config/moodle/custom.scss`", "All portal styling"],
       ["`native/stack/moodle/local/sebkiosk/exam-ui.js`", "Front-end runtime behaviour"],
       ["`scripts/moodle/post_install.php`", "First-run configuration \u2014 roles, SEB defaults, proctoring, theme"]],
      widths=[3.3, 3.0])
h3("Accounts")
table(["Account", "Role", "Credentials"],
      [["`examadmin`", "Site administrator", "In `.env`, held by the exam cell"],
       ["`invig01`", "Exam Hall Invigilator (demo)", "Issued separately"],
       ["`s2026\u2026`", "Student (test roster)", "Issued separately"]],
      widths=[1.3, 2.4, 2.6])
p = doc.add_paragraph()
r = p.add_run("Credentials are not printed in this document. Test and demo passwords are issued by "
              "the exam cell for the duration of a demo only.")
r.font.size = Pt(9); r.font.color.rgb = SLATE; r.italic = True

footer_text()

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("written:", OUT)
